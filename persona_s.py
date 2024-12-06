import os
import json
import streamlit as st
from docx import Document
from io import BytesIO
from PyPDF2 import PdfFileWriter, PdfFileReader

# Function to convert JSON chat history to Word document
def convert_json_to_word(json_data):
    doc = Document()
    
    for message in json_data:
        role = message['role']
        content = message['content']
        
        # Add role and content to the document
        p = doc.add_paragraph()
        p.add_run(f"{role.capitalize()}: ").bold = True
        p.add_run(content)
    
    return doc

# Function to save the Word document with password protection
def save_word_with_password(doc, password):
    # Save the document to a BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    
    # Convert the Word document to PDF with password protection
    pdf_output_path = 'chat_history.pdf'
    with open(pdf_output_path, 'wb') as f:
        f.write(byte_io.getvalue())
    
    # Add password protection to the PDF
    pdf_writer = PdfFileWriter()
    pdf_reader = PdfFileReader(pdf_output_path)
    
    for page_num in range(pdf_reader.numPages):
        pdf_writer.addPage(pdf_reader.getPage(page_num))
    
    pdf_writer.encrypt(password)
    
    protected_pdf_output_path = 'protected_chat_history.pdf'
    with open(protected_pdf_output_path, 'wb') as f:
        pdf_writer.write(f)
    
    return protected_pdf_output_path

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

# Define the characters with their segments
characters = {
    "Balance Seekers": {
        "Name": "Ben",
        "Demographics": "Typically middle-aged, balanced gender distribution, moderate income.",
        "Personality": "Balanced, health-conscious, moderate.",
        "Health Status": "Generally healthy.",
        "Provider Utilization": "Moderate.",
        "Insurance": "Typically insured.",
        "Health Insurance Coverage": "Comprehensive.",
        "Chronic Conditions": "Few or none.",
        "Words/Phrases (recommended)": "Balance, wellness, moderation.",
        "Words to Lose": "Extreme, neglect."
    },
    "Priority Juggler": {
        "Name": "Paige",
        "Demographics": "Often parents, middle-aged, moderate to high income.",
        "Personality": "Busy, caring, multitasker.",
        "Health Status": "Varies.",
        "Provider Utilization": "Inconsistent.",
        "Insurance": "Varies.",
        "Health Insurance Coverage": "Varies.",
        "Chronic Conditions": "Possible.",
        "Words/Phrases (recommended)": "Manage, prioritize, balance.",
        "Words to Lose": "Neglect, ignore."
    },
    "Willful Endurer": {
        "Name": "Willy",
        "Demographics": "Older adults, often retired, varied income.",
        "Personality": "Resilient, determined, enduring.",
        "Health Status": "Chronic conditions.",
        "Provider Utilization": "High.",
        "Insurance": "Typically insured.",
        "Health Insurance Coverage": "Comprehensive.",
        "Chronic Conditions": "Multiple.",
        "Words/Phrases (recommended)": "Endure, resilience, strength.",
        "Words to Lose": "Weak, give up."
    },
    "Self Achiever": {
        "Name": "Sarah",
        "Demographics": "Young to middle-aged adults, high income, career-focused.",
        "Personality": "Goal-oriented, proactive, ambitious.",
        "Health Status": "Generally healthy.",
        "Provider Utilization": "Proactive.",
        "Insurance": "Typically insured.",
        "Health Insurance Coverage": "Comprehensive.",
        "Chronic Conditions": "Few or none.",
        "Words/Phrases (recommended)": "Achieve, proactive, success.",
        "Words to Lose": "Lazy, passive."
    },
    "Trustful Responder": {
        "Name": "Tim",
        "Demographics": "Varied age, typically insured, moderate to high income.",
        "Personality": "Trusting, compliant, responsive.",
        "Health Status": "Varies.",
        "Provider Utilization": "High.",
        "Insurance": "Typically insured.",
        "Health Insurance Coverage": "Comprehensive.",
        "Chronic Conditions": "Possible.",
        "Words/Phrases (recommended)": "Trust, follow, reliable.",
        "Words to Lose": "Skeptical, ignore."
    },
    "General Bot": {
        "Name": "Brandience",
        "Demographics": "No Age Restriction, varied Income",
        "Personality": "Open, honest, responsive.",
        "Health Status": "Varies.",
        "Provider Utilization": "High.",
        "Insurance": "Typically insured.",
        "Health Insurance Coverage": "Comprehensive.",
        "Chronic Conditions": "Not possible"
    }
}

if "messages" not in st.session_state:
    st.session_state["messages"] = [
        {"role": "system", "content": "You are a friendly agent where you answer all the questions also remember about your characteristics, do not forget the name and answer your health conditions based on that behavior."}
    ]

# Function to append the response to the chat history
def append_to_chat_history(response):
    st.session_state["messages"].append({
        "role": response.role,
        "content": response.content
    })

# Function to save the chat history to a file
def save_chat_history_to_file(filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(st.session_state["messages"], f, ensure_ascii=False)

# Password input
password_placeholder = st.empty()
password = password_placeholder.text_input("Enter Password:", type="password", key="password_input")

if password == 'Brandpersona123':  # Replace with your desired password
    password_placeholder.empty()  # Remove the password input field

    # Select a character
    character = st.sidebar.selectbox('Select a Character', list(characters.keys()))

    # Get the bot's name for the selected character
    bot_name = characters[character]['Name']

    st.title(f'🧑‍💻 {bot_name} Online 💬 Chatbot')
    st.write(f'My name is {bot_name}🤖. I know many things, ask me anything you like, but please, don\'t ask me stupid questions❓')

    # Display character description
    st.header(f'{character}')
    for key, value in characters[character].items():
       st.subheader(key)
       st.write(value)

    # General questions section
    st.subheader('Ask General Questions')
    general_question = st.text_input('Ask a general question about this character\'s behaviors or values:')

    # Get GroqCloud AI response for general questions
    if st.button('Ask General Question'):
       if general_question:
           st.session_state['messages'].append({'role': 'user', 'content': f'As a {character}, {general_question}'})
           st.chat_message('user').write(general_question)
           response = client.chat.completions.create(
               messages=st.session_state['messages'],
               model='llama3-8b-8192'
           )
           msg = response.choices[0].message
           append_to_chat_history(msg)
           st.chat_message('assistant').write(msg.content)
           save_chat_history_to_file('chat_history.json')
       else:
           st.write('Please enter a question.')

    # Advertising creative section
    st.subheader('Test Opinions on Advertising Creative')
    creative_input = st.text_input('Enter a headline or description of an image for the character to review:')

    # Get GroqCloud AI response for advertising creative
    if st.button('Test Creative'):
       if creative_input:
           st.session_state['messages'].append({'role': 'user', 'content': f'As a {character}, what do you think about this: {creative_input}'})
           st.chat_message('user').write(creative_input)
           response = client.chat.completions.create(
               messages=st.session_state['messages'],
               model='llama3-8b-8192'
           )
           msg = response.choices[0].message
           append_to_chat_history(msg)
           st.chat_message('assistant').write(msg.content)
           save_chat_history_to_file('chat_history.json')
       else:
           st.write('Please enter a headline or description.')

    # Button to download chat history as Word document with password protection
    if st.button("Download Chat History as Word Document"):
        with open('chat_history.json', 'r', encoding='utf-8') as f:
            chat_history_json = json.load(f)
        
        doc = convert_json_to_word(chat_history_json)
        
        password = st.text_input("Enter a password for the document:", type="password")
        
        if password:
            protected_pdf_path = save_word_with_password(doc, password)
            
            with open(protected_pdf_path, 'rb') as f:
                pdf_data = f.read()
            
            st.download_button(label="Download Protected Word Document", data=pdf_data, file_name="protected_chat_history.pdf", mime="application/pdf")
        else:
            st.write("Please enter a password.")
